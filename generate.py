import os
import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

def format_text(paragraph, text, bold=False, align=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    paragraph.alignment = align if align else WD_PARAGRAPH_ALIGNMENT.LEFT  # Ajusta o alinhamento corretamente


def add_file_to_doc(doc, dir_name, file_path, add_page_break=True):
    if add_page_break:
        doc.add_page_break()
    relative_path = os.path.splitext(os.path.relpath(file_path, dir_name))[0].replace(os.sep, '.')
    prefix_doc = "diretório.arquivo : "
    title_text = prefix_doc + relative_path
    format_text(doc.add_paragraph(title_text), "", bold=True)
    doc.add_paragraph("")
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
            doc.add_paragraph(file.read())
    except Exception as e:
        doc.add_paragraph(f"Erro ao ler o arquivo: {file_path}\n{str(e)}")

def generate_document(base_dir, output_filename, include_dirs, include_env=True):
    doc = Document()
    title = "Linguagem de Programação III - Entrega 1 - Guilherme Silva Sampaio"
    format_text(doc.add_paragraph(title), "", bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    doc.add_paragraph("")
    
    files_list = []
    for root, dirs, files in os.walk(base_dir):
        if "node_modules" in root:
            continue
        relative_path = os.path.relpath(root, base_dir)
        if any(relative_path.startswith(d) for d in include_dirs):
            for file in files:
                if file.endswith(".env") and not include_env:
                    continue
                file_path = os.path.join(root, file)
                files_list.append((root, file_path))
    
    files_list.sort(key=lambda x: x[1])
    
    first_file = True
    for root, file_path in files_list:
        if first_file:
            add_file_to_doc(doc, base_dir, file_path, add_page_break=False)
            first_file = False
        else:
            add_file_to_doc(doc, base_dir, file_path)
    
    doc.add_page_break()
    date_str = datetime.datetime.now().strftime("%d/%m/%Y")
    format_text(doc.add_paragraph(f"Dourados, {date_str} -- (Assinatura)"), "")
    
        # Criar diretório de saída se não existir
    output_dir = "arquivos-entrega"
    os.makedirs(output_dir, exist_ok=True)
    
    # Salvar arquivos no diretório de saída
    docx_path = os.path.join(output_dir, output_filename + ".docx")
    pdf_path = os.path.join(output_dir, output_filename + ".pdf")
    
    doc.save(docx_path)
    convert(docx_path, pdf_path)

# Gerar os documentos
base_dir_front = "front-end"
base_dir_back = "back-end"

generate_document(base_dir_front, "front-end", ["public", "src"], include_env=False)
generate_document(base_dir_back, "back-end", ["src"], include_env=True)


