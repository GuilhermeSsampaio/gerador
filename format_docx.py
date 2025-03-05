from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def create_document():
    """Cria e retorna um novo documento com configurações padrão."""
    return Document()

def add_title(doc, title):
    """Adiciona um título centralizado e em negrito."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Cria o run com o texto e aplica negrito
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    return paragraph

def add_subtitle(doc, subtitle):
    """Adiciona um subtítulo em negrito e alinhado à esquerda."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Cria o run com o texto e aplica negrito
    run = paragraph.add_run(subtitle)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    return paragraph

def add_paragraph_text(doc, text, bold=False, align=WD_PARAGRAPH_ALIGNMENT.LEFT):
    """Adiciona um parágrafo com formatação específica."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    
    # Cria o run com o texto
    run = paragraph.add_run(text)
    
    # Aplica negrito se solicitado
    if bold:
        run.bold = True
    
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    return paragraph

def add_empty_paragraph(doc):
    """Adiciona um parágrafo vazio."""
    return doc.add_paragraph()

def add_page_break(doc):
    """Adiciona uma quebra de página."""
    doc.add_page_break()

def save_document(doc, file_path):
    """Salva o documento no caminho especificado."""
    doc.save(file_path)

# Função para verificar se o negrito está funcionando (teste)
def test_bold_functionality(output_path="test_bold.docx"):
    doc = create_document()
    
    # Adiciona título em negrito
    add_title(doc, "Este é um título em negrito")
    
    # Adiciona subtítulo em negrito
    add_subtitle(doc, "Este é um subtítulo em negrito")
    
    # Adiciona texto normal
    add_paragraph_text(doc, "Este é um texto normal")
    
    # Adiciona texto em negrito
    add_paragraph_text(doc, "Este texto deve estar em negrito", bold=True)
    
    # Salva para testar
    save_document(doc, output_path)
    print(f"Documento de teste salvo como {output_path}")