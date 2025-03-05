from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt  # Importe Pt para definir tamanho de fonte

document = Document()

def add_title(title):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    # Aplica a fonte e tamanho aqui também
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)  # Tamanho para o título
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Título centralizado

def add_paragraph_with_formatting(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):  # Default LEFT
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    # Aplica a formatação de fonte padrão
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    # Aplica o alinhamento (padrão LEFT se não for especificado)
    paragraph.alignment = align
    
    return paragraph

def boldText(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    # Aplica a formatação de fonte padrão
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return run

def alignText(paragraph, align):
    paragraph.alignment = align

def set_document_default_alignment():
    # Define o alinhamento padrão para LEFT para todos os estilos de parágrafo no documento
    for style in document.styles:
        if hasattr(style, 'paragraph_format'):
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Define o alinhamento padrão para LEFT
set_document_default_alignment()

# Exemplo de uso
add_title("Linguagem de Programação III - Entrega 1 - Guilherme Silva Sampaio")

# Adicione um parágrafo formatado (com alinhamento LEFT por padrão)
p1 = add_paragraph_with_formatting("Este é um texto com formatação aplicada.", bold=False)

# Adicione texto em negrito ao parágrafo existente
p2 = document.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Define explicitamente como LEFT
boldText(p2, "Este texto está em negrito. ")
p2.add_run("Este texto não está em negrito, mas mantém a fonte Times New Roman.").font.name = 'Times New Roman'
p2.add_run(" Continuando o texto...").font.name = 'Times New Roman'

# Salva o documento
document.save("test.docx")